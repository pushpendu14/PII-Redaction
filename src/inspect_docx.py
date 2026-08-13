from docx import Document

document = Document("input/Red Herring Prospectus.docx")

print("========== DOCUMENT INFO ==========")
print("Paragraphs:", len(document.paragraphs))
print("Tables:", len(document.tables))

print("\n========== NON-EMPTY PARAGRAPHS ==========")

count = 0

for i, paragraph in enumerate(document.paragraphs):
    text = paragraph.text.strip()

    if text:
        print(f"[Paragraph {i}] {text}")
        count += 1

    if count >= 20:
        break

print("\n========== FIRST TABLE ==========")

if document.tables:
    table = document.tables[0]

    print("Rows:", len(table.rows))
    print("Columns:", len(table.columns))

    for row in table.rows[:5]:
        print([cell.text for cell in row.cells])
else:
    print("No tables found.")