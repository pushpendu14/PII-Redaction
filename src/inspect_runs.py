from docx import Document

document = Document("input/Red Herring Prospectus.docx")

print("========== PARAGRAPH RUN INSPECTION ==========")

checked = 0

for i, paragraph in enumerate(document.paragraphs):
    if not paragraph.text.strip():
        continue

    print(f"\nParagraph {i}:")
    print("Full text:", repr(paragraph.text))
    print("Number of runs:", len(paragraph.runs))

    for j, run in enumerate(paragraph.runs):
        print(f"  Run {j}: {repr(run.text)}")

    checked += 1

    if checked >= 10:
        break


print("\n========== TABLE CELL RUN INSPECTION ==========")

if document.tables:
    table = document.tables[0]

    for row_index, row in enumerate(table.rows[:2]):
        for col_index, cell in enumerate(row.cells[:3]):
            print(f"\nCell [{row_index}][{col_index}]:")
            print("Full text:", repr(cell.text))
            print("Number of paragraphs:", len(cell.paragraphs))

            for paragraph in cell.paragraphs:
                for j, run in enumerate(paragraph.runs):
                    print(f"  Run {j}: {repr(run.text)}")