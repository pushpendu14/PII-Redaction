from docx import Document

document = Document("input/Red Herring Prospectus.docx")

print("Document loaded successfully!")
print("Number of paragraphs:", len(document.paragraphs))

print("\nFirst 10 paragraphs:\n")

for paragraph in document.paragraphs[:10]:
    print(paragraph.text)